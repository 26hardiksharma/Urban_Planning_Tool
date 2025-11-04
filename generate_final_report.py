from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_bold_paragraph(label, text):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    return p

title = doc.add_heading('NAGPUR CITY PLANNER: OPTIMIZATION DASHBOARD', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Design and Analysis of Algorithms - Lab Project Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True

doc.add_paragraph()

add_heading('1. PROJECT OVERVIEW', 1)
add_bold_paragraph('Project Title: ', 'Nagpur City Planner: Optimization Dashboard')
add_bold_paragraph('Domain: ', 'Urban Planning and Smart City Management')
add_bold_paragraph('Technology Stack: ', 'Python, Flask, NetworkX, NumPy, SciPy, HTML/CSS, JavaScript')
add_bold_paragraph('Date: ', datetime.now().strftime('%B %Y'))

doc.add_paragraph()
p = doc.add_paragraph(
    'This project implements a comprehensive web-based optimization dashboard that applies '
    'six fundamental algorithms from Design and Analysis of Algorithms to solve real-world '
    'urban planning problems in the Nagpur metropolitan area. The application demonstrates '
    'practical implementations of graph theory, dynamic programming, computational geometry, '
    'and backtracking techniques to optimize resource allocation, logistics, and infrastructure planning.'
)

doc.add_page_break()
add_heading('2. SYSTEM ARCHITECTURE', 1)
add_heading('2.1 Project Structure', 2)

structure_text = """Urban_Planning_Tool/
├── app.py (Main Flask application)
├── requirements.txt (Python dependencies)
├── core_algorithms/
│   ├── shortestPath.py (Dijkstra's Algorithm)
│   ├── knapsackSolver.py (0/1 Knapsack)
│   ├── mstBuilder.py (Minimum Spanning Tree)
│   ├── convexHullBuilder.py (Convex Hull)
│   ├── nQueenPlacement.py (N-Queen)
│   └── closestPairSolver.py (Closest Pair)
├── data/
│   ├── city_network.json (50 nodes, 60 edges)
│   ├── projects.json (~30 urban projects)
│   └── coordinates.json (50 Nagpur localities)
├── templates/index.html
└── static/style.css"""

doc.add_paragraph(structure_text)

add_heading('2.2 Technology Stack', 2)
tech_table = doc.add_table(rows=5, cols=2)
tech_table.style = 'Light Grid Accent 1'
tech_data = [
    ['Backend Framework', 'Flask (Python)'],
    ['Graph Processing', 'NetworkX'],
    ['Numerical Computing', 'NumPy, SciPy'],
    ['Frontend', 'HTML5, CSS3, JavaScript'],
    ['Data Format', 'JSON']
]
for i, (tech, desc) in enumerate(tech_data):
    tech_table.rows[i].cells[0].text = tech
    tech_table.rows[i].cells[1].text = desc

doc.add_page_break()
add_heading('3. ALGORITHM IMPLEMENTATIONS', 1)

algorithms = [
    {'number': '1', 'name': 'Emergency Response Route Planner', 'algorithm': "Dijkstra's Shortest Path", 
     'complexity': 'O(E log V)', 'input': 'Start/End node IDs', 'output': 'Optimal path, distance',
     'use_case': 'Ambulance routing, emergency dispatch'},
    {'number': '2', 'name': 'Budget Maximizer', 'algorithm': '0/1 Knapsack (DP)', 
     'complexity': 'O(n x W)', 'input': 'Budget in Lakhs', 'output': 'Selected projects, benefit, cost',
     'use_case': 'Budget allocation, project prioritization'},
    {'number': '3', 'name': 'Utility Network Design', 'algorithm': 'Minimum Spanning Tree', 
     'complexity': 'O(E log V)', 'input': 'City network graph', 'output': 'Required connections, cost',
     'use_case': 'Utility grids, infrastructure planning'},
    {'number': '4', 'name': 'Define Land Boundary', 'algorithm': 'Convex Hull', 
     'complexity': 'O(n log n)', 'input': '50 locality coordinates', 'output': 'Boundary area, vertices',
     'use_case': 'Zoning, land acquisition, boundary definition'},
    {'number': '5', 'name': 'Non-Interfering Sensor Placement', 'algorithm': 'N-Queen Backtracking', 
     'complexity': 'O(N!)', 'input': 'Number of sensors (1-14)', 'output': 'Coordinates, visual grid',
     'use_case': 'IoT deployment, wireless networks'},
    {'number': '6', 'name': 'Emergency Proximity Checker', 'algorithm': 'Closest Pair', 
     'complexity': 'O(n^2)', 'input': '50 incident coordinates', 'output': 'Closest pair, distance',
     'use_case': 'Resource allocation, coverage gaps'}
]

for algo in algorithms:
    add_heading(f"3.{algo['number']} {algo['name']}", 2)
    add_bold_paragraph('Algorithm: ', algo['algorithm'])
    add_bold_paragraph('Time Complexity: ', algo['complexity'])
    add_bold_paragraph('Input: ', algo['input'])
    add_bold_paragraph('Output: ', algo['output'])
    add_bold_paragraph('Use Case: ', algo['use_case'])
    doc.add_paragraph()

doc.add_page_break()
add_heading('4. ALGORITHM COMPLEXITY ANALYSIS', 1)

analysis_table = doc.add_table(rows=7, cols=4)
analysis_table.style = 'Light Grid Accent 1'
headers = ['Feature', 'Algorithm', 'Time Complexity', 'Space']
for i, h in enumerate(headers):
    analysis_table.rows[0].cells[i].text = h
    analysis_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

analysis_data = [
    ['1', "Dijkstra's", 'O(E log V)', 'O(V)'],
    ['2', '0/1 Knapsack', 'O(n x W)', 'O(n x W)'],
    ['3', 'MST', 'O(E log V)', 'O(V)'],
    ['4', 'Convex Hull', 'O(n log n)', 'O(n)'],
    ['5', 'N-Queen', 'O(N!)', 'O(N)'],
    ['6', 'Closest Pair', 'O(n^2)', 'O(n)']
]
for i, data in enumerate(analysis_data, 1):
    for j, value in enumerate(data):
        analysis_table.rows[i].cells[j].text = value

doc.add_page_break()
add_heading('5. KEY FEATURES', 1)

features = [
    'Interactive web dashboard with 6 optimization algorithms',
    'Real-time algorithm execution and visualization',
    'Anchor-based navigation (no scroll-to-top)',
    'Network map: 50 nodes, 60 edges',
    'Form validation and error handling',
    'Visual N-Queen grid display',
    'Convex hull boundary visualization',
    'Formatted path and coordinate outputs'
]
for f in features:
    doc.add_paragraph(f'* {f}', style='List Bullet')

add_heading('6. TESTING AND VALIDATION', 1)
doc.add_paragraph('Tested with: various node pairs, budget ranges 0-1000L, full network MST, '
                  '50-point convex hull, N=1-14 sensors, 50 localities proximity, edge cases.')

doc.add_page_break()
add_heading('7. CHALLENGES AND SOLUTIONS', 1)

challenges = [
    ('Challenge', 'Solution'),
    ('Page scroll-to-top', 'Anchor links (#feature-X)'),
    ('N-Queen performance', 'Limited to N<=14'),
    ('Result display consistency', 'Centralized context rendering'),
    ('Convex hull edge cases', 'Minimum 3 points validation')
]
challenge_table = doc.add_table(rows=len(challenges), cols=2)
challenge_table.style = 'Light Grid Accent 1'
for i, (c, s) in enumerate(challenges):
    challenge_table.rows[i].cells[0].text = c
    challenge_table.rows[i].cells[1].text = s
    if i == 0:
        challenge_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        challenge_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

doc.add_page_break()
add_heading('8. FUTURE ENHANCEMENTS', 1)

enhancements = [
    'O(n log n) divide-and-conquer closest pair',
    'Interactive graph visualization (D3.js)',
    'Additional algorithms: Floyd-Warshall, A*',
    'Database support (PostgreSQL)',
    'Admin panel for data management',
    'Export to PDF/Excel',
    'Real-time animation',
    'Performance benchmarking'
]
for e in enhancements:
    doc.add_paragraph(f'* {e}', style='List Bullet')

doc.add_page_break()
add_heading('9. CONCLUSION', 1)

conclusion = """The Nagpur City Planner successfully demonstrates 6 DAA algorithms for urban planning:
graph algorithms (Dijkstra's, MST), dynamic programming (Knapsack), computational geometry
(Convex Hull), backtracking (N-Queen), and geometric algorithms (Closest Pair). The modular
architecture and web interface make complex algorithms accessible for real-world applications."""

doc.add_paragraph(conclusion)

add_heading('10. REFERENCES', 1)
refs = ['CLRS - Introduction to Algorithms', 'NetworkX Documentation', 'Flask Documentation',
        'SciPy Documentation', 'Algorithm Design Manual - Skiena']
for r in refs:
    doc.add_paragraph(f'* {r}', style='List Bullet')

doc.add_paragraph()
footer = doc.add_paragraph('--- End of Report ---')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].italic = True

output_path = r'd:\Urban_Planning_Tool\Nagpur_City_Planner_Report_2025.docx'
doc.save(output_path)
print(f"Report generated successfully!")
print(f"Location: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
