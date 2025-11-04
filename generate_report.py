from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_heading(text, level=1):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_bold_paragraph(label, text):
    p = doc.add_paragraph()
    p.add_run(label).bold = True
    p.add_run(text)
    return p

title = doc.add_heading('NAGPUR CITY PLANNER: OPTIMIZATION DASHBOARD', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Design and Analysis of Algorithms - Lab Project Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True

doc.add_paragraph()

add_heading('1. PROJECT OVERVIEW', 1)

add_bold_paragraph('Project Title: ', 'Nagpur City Planner: Optimization Dashboard')
add_bold_paragraph('Domain: ', 'Urban Planning and Smart City Management')
add_bold_paragraph('Technology Stack: ', 'Python, Flask, NetworkX, NumPy, HTML/CSS, JavaScript')
add_bold_paragraph('Date: ', datetime.now().strftime('%B %Y'))

doc.add_paragraph()

p = doc.add_paragraph(
    'This project implements a comprehensive web-based optimization dashboard that applies '
    'six fundamental algorithms from Design and Analysis of Algorithms to solve real-world '
    'urban planning problems in the Nagpur metropolitan area. The application demonstrates '
    'practical implementations of graph theory, dynamic programming, backtracking, and '
    'divide-and-conquer techniques to optimize resource allocation, logistics, and infrastructure planning.'
)

doc.add_page_break()

add_heading('2. SYSTEM ARCHITECTURE', 1)

add_heading('2.1 Project Structure', 2)

structure_text = """
Urban_Planning_Tool/
├── app.py                          # Main Flask application
├── requirements.txt                # Python dependencies
├── core_algorithms/                # Algorithm implementations
│   ├── shortestPath.py            # Dijkstra's Algorithm
│   ├── knapsackSolver.py          # 0/1 Knapsack (Dynamic Programming)
│   ├── mstBuilder.py              # Minimum Spanning Tree (Prim's/Kruskal's)
│   ├── multistageAnalyzer.py      # Multistage Graph Optimization
│   ├── nQueenPlacement.py         # N-Queen Backtracking
│   └── closestPairSolver.py       # Closest Pair Algorithm
├── data/                           # JSON data files
│   ├── city_network.json          # 50 nodes, 60 edges network
│   ├── projects.json              # ~30 urban projects
│   ├── coordinates.json           # 50 Nagpur localities
│   └── multistage_project.json    # Project task network
├── templates/
│   └── index.html                 # Main dashboard UI
└── static/
    └── style.css                  # Styling
"""

code_para = doc.add_paragraph(structure_text)
code_para.style = 'Normal'
for run in code_para.runs:
    run.font.name = 'Courier New'
    run.font.size = Pt(9)

add_heading('2.2 Technology Stack', 2)

tech_table = doc.add_table(rows=5, cols=2)
tech_table.style = 'Light Grid Accent 1'

tech_data = [
    ['Backend Framework', 'Flask (Python)'],
    ['Graph Processing', 'NetworkX'],
    ['Numerical Computing', 'NumPy, SciPy'],
    ['Frontend', 'HTML5, CSS3, JavaScript'],
    ['Data Format', 'JSON']
]

for i, (tech, desc) in enumerate(tech_data):
    tech_table.rows[i].cells[0].text = tech
    tech_table.rows[i].cells[1].text = desc

doc.add_page_break()

add_heading('3. ALGORITHM IMPLEMENTATIONS', 1)

algorithms = [
    {
        'number': '1',
        'name': 'Emergency Response Route Planner',
        'algorithm': "Dijkstra's Shortest Path Algorithm",
        'complexity': 'O(E log V) using priority queue',
        'description': 'Finds the fastest path for emergency vehicles through the city network. Uses NetworkX implementation of Dijkstra\'s algorithm on a weighted graph with 50 nodes (N1-N50) representing Nagpur localities and 60 edges representing road connections.',
        'input': 'Start node ID, End node ID',
        'output': 'Optimal path sequence and total distance/time',
        'use_case': 'Optimizing ambulance routes, fire truck dispatch, police response times'
    },
    {
        'number': '2',
        'name': 'Budget Maximizer',
        'algorithm': '0/1 Knapsack (Dynamic Programming)',
        'complexity': 'O(n × W) where n = projects, W = budget',
        'description': 'Selects the optimal mix of urban development projects that maximizes public benefit while staying within budget constraints. Uses dynamic programming with a 2D table to compute maximum benefit values.',
        'input': 'Maximum budget (in Lakhs)',
        'output': 'Selected projects, total benefit, total cost',
        'use_case': 'Annual budget allocation, prioritizing infrastructure projects'
    },
    {
        'number': '3',
        'name': 'Utility Network Design',
        'algorithm': 'Minimum Spanning Tree (MST)',
        'complexity': 'O(E log V) using Prim\'s or Kruskal\'s',
        'description': 'Designs the least expensive utility network (water, electricity, sewage) that connects all service points. Uses NetworkX MST implementation to find minimum total construction cost.',
        'input': 'City network graph',
        'output': 'Required connections and minimum total cost',
        'use_case': 'Planning utility grids, minimizing infrastructure costs'
    },
    {
        'number': '4',
        'name': 'Project Critical Path Analyzer',
        'algorithm': 'Multistage Graph Optimization',
        'complexity': 'O(V + E) for directed acyclic graph',
        'description': 'Finds the sequence of project tasks that results in minimum total cost or time. Treats the network as a multistage graph and uses shortest path algorithms to identify the critical path from N1 to N50.',
        'input': 'Project network with task dependencies',
        'output': 'Optimal task sequence and minimum total cost',
        'use_case': 'Project scheduling, construction planning, resource optimization'
    },
    {
        'number': '5',
        'name': 'Non-Interfering Sensor Placement',
        'algorithm': 'N-Queen Backtracking',
        'complexity': 'O(N!) in worst case',
        'description': 'Places N sensors/Wi-Fi hotspots on an N×N grid such that no two interfere with each other (no two on same row, column, or diagonal). Uses recursive backtracking to find valid placements.',
        'input': 'Number of sensors (N, limited to 1-14)',
        'output': 'Sensor coordinates, visual grid, total solutions found',
        'use_case': 'IoT sensor deployment, wireless network planning, surveillance optimization'
    },
    {
        'number': '6',
        'name': 'Emergency Response Proximity Checker',
        'algorithm': 'Closest Pair Algorithm',
        'complexity': 'O(n²) brute force implementation',
        'description': 'Identifies the two incident sites closest to each other among 50 Nagpur localities to optimize dual dispatch and resource sharing. Uses Euclidean distance calculation.',
        'input': '50 emergency incident coordinates',
        'output': 'Closest pair names, distance, and coordinates',
        'use_case': 'Emergency resource allocation, identifying coverage gaps'
    }
]

for algo in algorithms:
    add_heading(f"3.{algo['number']} {algo['name']}", 2)
    
    add_bold_paragraph('Algorithm: ', algo['algorithm'])
    add_bold_paragraph('Time Complexity: ', algo['complexity'])
    
    doc.add_paragraph()
    doc.add_paragraph(algo['description'])
    
    doc.add_paragraph()
    add_bold_paragraph('Input: ', algo['input'])
    add_bold_paragraph('Output: ', algo['output'])
    add_bold_paragraph('Real-World Use Case: ', algo['use_case'])
    
    doc.add_paragraph()

doc.add_page_break()

add_heading('4. DATA STRUCTURES', 1)

data_desc = [
    ('city_network.json', '50 nodes representing Nagpur localities (N1-N50), 60 weighted edges representing road connections. Used by Features 1, 3, and 4.'),
    ('projects.json', 'Approximately 30 urban development projects with cost (in Lakhs) and benefit values. Used by Feature 2.'),
    ('coordinates.json', '50 Nagpur locality coordinates (x, y) for spatial analysis. Used by Feature 6.'),
    ('multistage_project.json', 'Project task network with nodes and dependencies for critical path analysis. Used by Feature 4.')
]

for filename, desc in data_desc:
    add_bold_paragraph(f'{filename}: ', desc)

doc.add_page_break()

add_heading('5. KEY FEATURES', 1)

features = [
    'Interactive web-based dashboard with responsive design',
    'Real-time algorithm execution and result visualization',
    'Form-based input with validation',
    'Dynamic result display with formatted output',
    'Automatic scroll-to-section after form submission',
    'Network map display showing all nodes and connections',
    'Support for 50-node Nagpur city network',
    'Error handling for invalid inputs',
    'Visual grid display for sensor placement',
    'Formatted path and coordinate outputs'
]

for feature in features:
    doc.add_paragraph(f'• {feature}', style='List Bullet')

doc.add_paragraph()

add_heading('6. IMPLEMENTATION DETAILS', 1)

add_heading('6.1 Backend Implementation', 2)

backend_text = """
The Flask application (app.py) serves as the main controller:
• Route handlers process form submissions
• Data loading functions read JSON files
• Algorithm modules are imported and called with appropriate parameters
• Results are rendered back to the HTML template using Jinja2
• Centralized context rendering ensures consistent data availability
"""

doc.add_paragraph(backend_text)

add_heading('6.2 Frontend Implementation', 2)

frontend_text = """
The HTML/CSS interface provides:
• Structured sections for each algorithm (features 1-6)
• Form inputs with placeholders and validation
• Result boxes with color-coded success indicators
• Scrollable network map display
• Anchor-based navigation to prevent page-top scrolling
• Visual grids for sensor placement
• Formatted output for paths and coordinates
"""

doc.add_paragraph(frontend_text)

add_heading('6.3 Algorithm Modules', 2)

modules_text = """
Each algorithm is implemented in a separate Python module:
• shortestPath.py: Uses NetworkX's Dijkstra implementation
• knapsackSolver.py: NumPy-based dynamic programming solution
• mstBuilder.py: NetworkX MST implementation
• multistageAnalyzer.py: Graph-based shortest path on directed network
• nQueenPlacement.py: Recursive backtracking with safety checks
• closestPairSolver.py: Brute-force O(n²) distance comparison
"""

doc.add_paragraph(modules_text)

doc.add_page_break()

add_heading('7. ALGORITHM ANALYSIS', 1)

analysis_table = doc.add_table(rows=7, cols=4)
analysis_table.style = 'Light Grid Accent 1'

headers = ['Feature', 'Algorithm', 'Time Complexity', 'Space Complexity']
for i, header in enumerate(headers):
    cell = analysis_table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True

analysis_data = [
    ['1', "Dijkstra's", 'O(E log V)', 'O(V)'],
    ['2', '0/1 Knapsack', 'O(n × W)', 'O(n × W)'],
    ['3', 'MST', 'O(E log V)', 'O(V)'],
    ['4', 'Multistage Graph', 'O(V + E)', 'O(V)'],
    ['5', 'N-Queen', 'O(N!)', 'O(N)'],
    ['6', 'Closest Pair', 'O(n²)', 'O(n)']
]

for i, data in enumerate(analysis_data, 1):
    for j, value in enumerate(data):
        analysis_table.rows[i].cells[j].text = value

doc.add_paragraph()

add_heading('8. TESTING AND VALIDATION', 1)

testing_text = """
The application has been tested with:
• Various node combinations for shortest path (N1-N50)
• Different budget constraints for knapsack (0-1000 Lakhs)
• Full network MST computation
• Multistage path from N1 to N50
• Sensor placement for N=1 to N=14
• Proximity checking on all 50 localities
• Edge cases: invalid nodes, zero budget, impossible placements
• Performance testing with maximum input sizes
"""

doc.add_paragraph(testing_text)

doc.add_page_break()

add_heading('9. CHALLENGES AND SOLUTIONS', 1)

challenges = [
    ('Challenge', 'Solution'),
    ('Page scrolling to top after form submission', 'Added anchor links (#feature-X) to form actions'),
    ('Complex multistage graph structure', 'Simplified to use Dijkstra\'s on full network'),
    ('N-Queen performance for large N', 'Limited input to N≤14 for stable computation'),
    ('Coordinate data structure consistency', 'Standardized JSON format with name/x/y fields'),
    ('Result display across multiple features', 'Centralized context rendering function')
]

challenge_table = doc.add_table(rows=len(challenges), cols=2)
challenge_table.style = 'Light Grid Accent 1'

for i, (challenge, solution) in enumerate(challenges):
    challenge_table.rows[i].cells[0].text = challenge
    challenge_table.rows[i].cells[1].text = solution
    if i == 0:
        challenge_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        challenge_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

add_heading('10. FUTURE ENHANCEMENTS', 1)

enhancements = [
    'Implement optimized O(n log n) closest pair algorithm using divide-and-conquer',
    'Add visualization of paths on interactive network graphs',
    'Implement additional algorithms: Floyd-Warshall, Bellman-Ford, A*',
    'Add database support for dynamic data management',
    'Create admin panel for data editing',
    'Export results to PDF/Excel',
    'Add user authentication and saved sessions',
    'Implement real-time algorithm execution animation',
    'Mobile-responsive design improvements',
    'Performance benchmarking dashboard'
]

for enhancement in enhancements:
    doc.add_paragraph(f'• {enhancement}', style='List Bullet')

doc.add_page_break()

add_heading('11. CONCLUSION', 1)

conclusion = """
The Nagpur City Planner Optimization Dashboard successfully demonstrates the practical application 
of six fundamental algorithms from Design and Analysis of Algorithms in solving real-world urban 
planning challenges. The project showcases:

1. Comprehensive implementation of graph algorithms (Dijkstra's, MST, Multistage Graph)
2. Dynamic programming solution for resource optimization (0/1 Knapsack)
3. Backtracking technique for constraint satisfaction (N-Queen)
4. Geometric algorithm for spatial analysis (Closest Pair)

The web-based interface makes these complex algorithms accessible and interactive, allowing users 
to solve actual urban planning problems with intuitive inputs and clear visualizations. The modular 
architecture ensures maintainability and extensibility for future enhancements.

This project bridges the gap between theoretical algorithm knowledge and practical software 
engineering, demonstrating how algorithmic thinking can optimize decision-making in smart city 
management and urban infrastructure planning.
"""

doc.add_paragraph(conclusion)

doc.add_paragraph()

add_heading('12. REFERENCES', 1)

references = [
    'Introduction to Algorithms (CLRS) - Thomas H. Cormen, Charles E. Leiserson, Ronald L. Rivest, Clifford Stein',
    'NetworkX Documentation - https://networkx.org/',
    'Flask Web Framework Documentation - https://flask.palletsprojects.com/',
    'Algorithm Design Manual - Steven S. Skiena',
    'NumPy Documentation - https://numpy.org/doc/'
]

for ref in references:
    doc.add_paragraph(f'• {ref}', style='List Bullet')

doc.add_paragraph()
doc.add_paragraph()

footer_para = doc.add_paragraph('--- End of Report ---')
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_para.runs[0].italic = True

output_path = 'Nagpur_City_Planner_Project_Report.docx'
doc.save(output_path)
print(f"Report generated successfully: {output_path}")
